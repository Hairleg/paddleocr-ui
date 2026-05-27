"""
Author: sizhchan
Org: dgaudit
Version: v0.1
Date: 2026-05-27
"""

#!/usr/bin/env python3
"""Compare our MinerU pipeline DOCX vs reference commercial OCR DOCX for all 3 files."""
from docx import Document
import os

cases = [
    ("江苏南通合同1", "output/江苏南通合同1_mineru.docx", "example/江苏南通合同1.docx"),
    ("江苏南通合同2", "output/江苏南通合同2_mineru.docx", "example/江苏南通合同2.docx"),
    ("浙江温岭合同", "output/浙江温岭合同_mineru.docx", "example/浙江温岭合同.docx"),
]

for name, our_path, ref_path in cases:
    if not os.path.exists(our_path):
        print(f"\n{'='*60}\n{name}: MISSING {our_path}\n{'='*60}")
        continue

    our = Document(our_path)
    ref = Document(ref_path)

    print(f"\n{'='*60}")
    print(f"{name}")
    print(f"{'='*60}")

    # Paragraph count
    our_paras = [p.text.strip() for p in our.paragraphs if p.text.strip()]
    ref_paras = [p.text.strip() for p in ref.paragraphs if p.text.strip()]
    print(f"Paragraphs:  our={len(our_paras)}  ref={len(ref_paras)}")

    # Table count
    print(f"Tables:      our={len(our.tables)}  ref={len(ref.tables)}")

    # Table header comparison
    our_headers = set()
    for t in our.tables:
        if t.rows:
            h = tuple(c.text.strip()[:30] for c in t.rows[0].cells if c.text.strip())
            if h: our_headers.add(h)
    ref_headers = set()
    for t in ref.tables:
        if t.rows:
            h = tuple(c.text.strip()[:30] for c in t.rows[0].cells if c.text.strip())
            if h: ref_headers.add(h)

    matched_headers = our_headers & ref_headers
    print(f"Table headers matched: {len(matched_headers)}/{len(ref_headers)}")

    # Missing tables (ref headers not in ours)
    missing = ref_headers - our_headers
    if missing:
        print(f"Missing table headers:")
        for h in missing:
            print(f"  {list(h)[:4]}")

    # Extra tables (ours not in ref)
    extra = our_headers - ref_headers
    if extra:
        print(f"Extra table headers:")
        for h in extra:
            print(f"  {list(h)[:4]}")

    # Text coverage: how many reference paragraphs appear in our output
    our_text_set = set(our_paras)
    ref_text_set = set(ref_paras)
    covered = 0
    for rt in ref_text_set:
        # Fuzzy match: check if first 20 chars appear
        prefix = rt[:20]
        for ot in our_text_set:
            if prefix and prefix in ot:
                covered += 1
                break
    print(f"Text coverage: {covered}/{len(ref_text_set)} ({100*covered/max(1,len(ref_text_set)):.0f}%)")

    # Show first 3 table headers from each
    print("\n  Our table headers (first 3):")
    for i, t in enumerate(our.tables[:3]):
        if t.rows:
            h = [c.text.strip()[:25] for c in t.rows[0].cells if c.text.strip()]
            print(f"    Table {i+1}: {h}")
    print("  Ref table headers (first 3):")
    for i, t in enumerate(ref.tables[:3]):
        if t.rows:
            h = [c.text.strip()[:25] for c in t.rows[0].cells if c.text.strip()]
            print(f"    Table {i+1}: {h}")
