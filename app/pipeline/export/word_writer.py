"""
Author: sizhchan
Org: dgaudit
Version: v0.2.0
Date: 2026-06-01
"""

"""
Word (.docx) export module.

Converts DocumentLayout IR into a python-docx Document,
preserving reading order, page breaks, and basic formatting.
"""

import logging

from docx import Document as DocxDocument
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.pipeline.ir.types import (
    DocumentLayout,
    ElementType,
    PageElement,
    PageLayout,
    PageSource,
    TextSpan,
)
from app.pipeline.export.font_rules import apply_government_font_rules

logger = logging.getLogger(__name__)

# Default Chinese font for paragraphs when source font is unknown
DEFAULT_FONT_NAME = "SimSun"
DEFAULT_FONT_SIZE = Pt(12)

# Mapping of common font names detected in PDFs to available fonts
FONT_FALLBACK_MAP = {
    "SimSun": "SimSun",
    "宋体": "SimSun",
    "SimHei": "SimHei",
    "黑体": "SimHei",
    "KaiTi": "KaiTi",
    "楷体": "KaiTi",
    "FangSong": "FangSong",
    "仿宋": "FangSong",
    "FZXiaoBiaoSong": "SimSun",  # Fallback: 方正小标宋 -> 宋体
    "方正小标宋简体": "SimSun",
}




def _merge_adjacent_paragraphs(elements):
    """Merge adjacent text elements with similar font sizes into flowing paragraphs.
    
    保守策略 (docx):
    - 字号差异 > 1.5pt → 不合并
    - 垂直间距 > 2×字号 → 不合并 (跨行/跨区域)
    """
    if len(elements) < 2:
        return
    merged = []
    prev = None
    for elem in elements:
        if elem.type != ElementType.PARAGRAPH or elem.type != getattr(prev, 'type', None):
            merged.append(elem)
            prev = elem
            continue
        if not prev or not hasattr(prev, 'content') or not prev.content:
            merged.append(elem)
            prev = elem
            continue
        if not elem.content:
            merged.append(elem)
            prev = elem
            continue

        # Font size check (1.5pt tolerance for docx)
        prev_size = prev.content[-1].font_size if prev.content else 12
        cur_size = elem.content[0].font_size if elem.content else 12
        if abs(prev_size - cur_size) >= 1.5:
            merged.append(elem)
            prev = elem
            continue

        # Position check: vertical gap must be reasonable
        px1, py1, pw1, ph1 = prev.bbox
        px2, py2, pw2, ph2 = elem.bbox
        line_height = max(ph1, prev_size * 1.5)
        vert_gap = abs(py2 - py1)
        if vert_gap > line_height * 2:
            merged.append(elem)
            prev = elem
            continue

        # Merge
        prev.content.extend(elem.content)
        prev.bbox = (min(px1, px2), min(py1, py2),
                     max(px1+pw1, px2+pw2)-min(px1, px2),
                     max(py1+ph1, py2+ph2)-min(py1, py2))
    elements[:] = merged
def write_word(doc_layout: DocumentLayout, output_path: str) -> str:
    """
    Write DocumentLayout to a .docx file.

    Args:
        doc_layout: The parsed document structure.
        output_path: Path for the output .docx file.

    Returns:
        The output path (confirmed written).
    """
    doc = DocxDocument()

    # Set default font for the document
    style = doc.styles["Normal"]
    font = style.font
    font.name = DEFAULT_FONT_NAME
    font.size = DEFAULT_FONT_SIZE

    for page_idx, page in enumerate(doc_layout.pages):
        # Apply government document font rules if applicable
        apply_government_font_rules(page.elements, page.height)
        # Merge adjacent paragraphs with similar font sizes (Chinese gov doc style)
        _merge_adjacent_paragraphs(page.elements)

        # Add page break between pages (not before the first)
        if page_idx > 0:
            doc.add_page_break()

        # Page marker disabled — cleaner output for production
        # _add_page_marker(doc, page)

        # Process elements in reading order (or all if no order specified)
        if page.reading_order:
            for elem_idx in page.reading_order:
                if elem_idx >= len(page.elements):
                    continue
                elem = page.elements[elem_idx]
                _write_element(doc, elem)
        else:
            for elem in page.elements:
                _write_element(doc, elem)

    doc.save(output_path)
    logger.info("Word document written: %s", output_path)
    return output_path


def _add_page_marker(doc: DocxDocument, page: PageLayout) -> None:
    """
    Add a small page indicator paragraph.
    This helps users orient themselves in the document.
    Can be removed later if undesired.
    """
    marker = doc.add_paragraph()
    marker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = marker.add_run(f"--- 第 {page.page_num} 页 ---")
    run.font.size = Pt(8)
    run.font.color.rgb = None  # Keep as-is (default dark)
    run.font.name = DEFAULT_FONT_NAME
    # Set the marker font to light gray for subtlety
    from docx.shared import RGBColor
    run.font.color.rgb = RGBColor(180, 180, 180)


def _write_element(doc: DocxDocument, elem: PageElement) -> None:
    """
    Write a single PageElement into the Word document.

    Dispatches based on element type.
    """
    if elem.type == ElementType.PARAGRAPH:
        _write_paragraph(doc, elem)
    elif elem.type == ElementType.IMAGE or elem.type == ElementType.STAMP:
        _write_image(doc, elem)
    elif elem.type == ElementType.TABLE:
        _write_table(doc, elem)
    # FORMULA type deferred to OMML injector module


def _write_paragraph(doc: DocxDocument, elem: PageElement) -> None:
    """Write a text paragraph element with inferred formatting."""
    if not elem.content or not isinstance(elem.content, list):
        return

    para = doc.add_paragraph()

    # Apply paragraph-level alignment based on first span's position
    x, y, w, h = elem.bbox
    page_center = 300  # approximate A4 center pt
    if x < 50:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif x > page_center + 50:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT  # default

    for span in elem.content:
        if not isinstance(span, TextSpan):
            continue
        run = para.add_run(span.text)
        _apply_span_formatting(run, span)


def _apply_span_formatting(run, span: TextSpan) -> None:
    """Apply font name, size, bold, and color from a TextSpan to a docx Run."""
    # Font name (with CJK font name support)
    font_name = span.font_name or DEFAULT_FONT_NAME
    mapped_font = FONT_FALLBACK_MAP.get(font_name, font_name)
    run.font.name = mapped_font
    # Set East Asian font name for CJK text
    from docx.oxml.ns import qn
    run._element.rPr.rFonts.set(qn('w:eastAsia'), mapped_font)

    # Font size
    if span.font_size:
        sz = span.font_size
        if sz > 22: sz = 16
        elif sz > 18: sz = 18
        run.font.size = Pt(sz)
    else:
        run.font.size = DEFAULT_FONT_SIZE

    # Bold
    if span.is_bold:
        run.bold = True

    # Color (if available)
    if hasattr(span, 'color') and span.color:
        try:
            from docx.shared import RGBColor
            hex_color = span.color.lstrip('#')
            if len(hex_color) == 6:
                run.font.color.rgb = RGBColor(
                    int(hex_color[0:2], 16),
                    int(hex_color[2:4], 16),
                    int(hex_color[4:6], 16),
                )
        except Exception:
            pass


def _write_image(doc: DocxDocument, elem: PageElement) -> None:
    """
    Embed an image element into the Word document.

    Images are placed inline (in reading order). The bbox from IR is
    used to scale the image proportionally and to align it (left/center/right)
    based on its horizontal position on the page.

    For stamps, images are typically centered and sized to ~2 inches wide.
    For photos, images are centered and scaled to fit page width.
    """
    if not elem.image_path:
        return

    try:
        para = doc.add_paragraph()

        # Use bbox X position to determine alignment
        x, y, w, h = elem.bbox
        page_center = 300  # approximate A4 center in pt
        if x < 50:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif x > page_center + 50:
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = para.add_run()

        # Scale stamp images smaller (they're usually ~2 inches)
        if elem.type == ElementType.STAMP:
            run.add_picture(elem.image_path, width=Inches(2.0))
        else:
            # Scale photo to fit within page width
            run.add_picture(elem.image_path, width=Inches(5.5))
    except FileNotFoundError:
        logger.warning("Image file not found for embedding: %s", elem.image_path)
    except Exception as exc:
        logger.warning("Failed to embed image %s: %s", elem.image_path, exc)


def _write_table(doc: DocxDocument, elem: PageElement) -> None:
    """
    Write a table element into the Word document.

    Uses TableData stored in elem._table_data (set by pipeline)
    for accurate grid structure with merged cells.
    """
    table_data = getattr(elem, '_table_data', None)

    if table_data and table_data.rows:
        rows = table_data.rows
        max_cols = max(len(row) for row in rows) if rows else 0
        if max_cols == 0:
            return

        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = "Table Grid"

        for r_idx, row_cells in enumerate(rows):
            for c_idx, cell_data in enumerate(row_cells):
                if c_idx >= max_cols:
                    break
                if cell_data.colspan == 0 or cell_data.rowspan == 0:
                    continue
                cell = table.cell(r_idx, c_idx)
                cell.text = cell_data.text

    elif elem.content and isinstance(elem.content, list) and elem.content:
        # Fallback: content is list of text lines
        rows = [[TableCell(row=i, col=0, text=str(s))] for i, s in enumerate(elem.content)]
        table = doc.add_table(rows=len(rows), cols=1)
        table.style = "Table Grid"
        for i, s in enumerate(elem.content):
            table.cell(i, 0).text = str(s)
